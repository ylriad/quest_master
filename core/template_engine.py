# core/template_engine.py
import os
from pathlib import Path
from datetime import datetime
from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML
from docx import Document
from docx.shared import Inches
import qrcode
from io import BytesIO

TEMPLATE_DIR = Path(__file__).parent.parent / "templates"
OUTPUT_DIR = Path("parchments")

class TemplateEngine:
    def __init__(self):
        self.env = Environment(loader=FileSystemLoader(TEMPLATE_DIR))

    def _ensure_output_dir(self):
        OUTPUT_DIR.mkdir(exist_ok=True)

    def _generate_qr_code(self, quest_id):
        """Generate QR code as base64 PNG for embedding in HTML."""
        qr = qrcode.QRCode(version=1, box_size=4, border=2)
        url = f"quest://id/{quest_id}"  # or real URL if available
        qr.add_data(url)
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="transparent")
        buffer = BytesIO()
        img.save(buffer, format="PNG")
        return "data:image/png;base64," + buffer.getvalue().decode("latin-1")

    def generate_pdf(self, quest_data, template_name="guild_contract.html"):
        self._ensure_output_dir()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{quest_data['id']}_{timestamp}.pdf"
        output_path = OUTPUT_DIR / filename

        # Add current date and QR (for bonus)
        context = {
            "quest": quest_data,
            "current_date": datetime.now().strftime("%d.%m.%Y"),
            "qr_code": self._generate_qr_code(quest_data["id"])
        }

        template = self.env.get_template(template_name)
        html_content = template.render(**context)
        HTML(string=html_content).write_pdf(str(output_path))
        return str(output_path)

    def generate_docx(self, quest_data):
        self._ensure_output_dir()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{quest_data['id']}_{timestamp}.docx"
        output_path = OUTPUT_DIR / filename

        doc = Document()
        doc.add_heading(f"📜 {quest_data['title']}", 0)
        doc.add_paragraph(f"ID: {quest_data['id']}")
        doc.add_paragraph(f"Сложность: {quest_data['difficulty']}")
        doc.add_paragraph(f"Награда: {quest_data['reward']} золотых")
        doc.add_paragraph(f"Дедлайн: {quest_data['deadline']}")
        doc.add_paragraph(f"Создано: {datetime.now().strftime('%d.%m.%Y')}")
        doc.add_heading("Описание", level=1)
        doc.add_paragraph(quest_data['description'])

        doc.save(str(output_path))
        return str(output_path)